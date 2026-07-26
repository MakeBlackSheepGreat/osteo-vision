"""Validate the generated Challenge Cup Word documents against their Markdown sources."""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, replace
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "research" / "reports" / "submission" / "challenge_cup_report_draft_20260721"

DATA_SUMMARY_THREE_MODEL_TERMS = (
    "残差注意力 U-Net（三种子均值）",
    "ConvNeXt 2D（单种子基线）",
    "多尺度深度可分离 U-Net（三种子均值）",
)

DATA_SUMMARY_THREE_MODEL_METRICS = (
    "0.914894",
    "0.843474",
    "0.909918",
    "0.004688",
    "5.131 ms",
    "26.032 MB",
    "0.898711",
    "0.816431",
    "0.890778",
    "0.003342",
    "3.566 ms",
    "22.193 MB",
    "0.897765",
    "0.814936",
    "0.893267",
    "0.006901",
    "4.254 ms",
    "20.139 MB",
)


@dataclass(frozen=True)
class DocumentSpec:
    name: str
    source: Path
    output: Path
    required_terms: tuple[str, ...]


SPECS = (
    DocumentSpec(
        name="完整报告",
        source=REPORT_DIR / "challenge_cup_feasibility_report_20260722_zh.md",
        output=REPORT_DIR / "颌骨骨髓炎智能化荧光诊疗比赛方案.docx",
        required_terms=("真实术中 ICG 颌骨骨髓炎", "医生复核", "82.026 ms"),
    ),
    DocumentSpec(
        name="精简提交报告",
        source=REPORT_DIR / "challenge_cup_concise_feasibility_report_20260726_zh.md",
        output=REPORT_DIR / "challenge_cup_concise_feasibility_report_20260726_zh_final.docx",
        required_terms=("15 份来源清单", "504 条记录", "医生复核"),
    ),
    DocumentSpec(
        name="数据汇总",
        source=REPORT_DIR / "数据汇总.md",
        output=REPORT_DIR / "平台软件数据汇总.docx",
        required_terms=(
            "3840×2160",
            "59.022 ms",
            "21.891 ms",
            "82.026 ms",
            "5.187 px → 1.361 px",
            "Dice",
            "0.917681",
            "IoU",
            "0.848335",
            "Recall",
            "0.9099",
            "0.724432 s",
            "5.776683 s",
            "723.579 MB",
            "0.036377 s",
            "0.176457 s",
            "380.134 MB",
            "0.010251 s",
            "0.047416 s",
            "109.454 MB",
            "真实术中 ICG 颌骨骨髓炎",
            *DATA_SUMMARY_THREE_MODEL_TERMS,
            *DATA_SUMMARY_THREE_MODEL_METRICS,
        ),
    ),
)

BANNED_INTERNAL_TERMS = (
    "keyframe_residual_attention_unet_s",
    "fluorescence_fusion_v2",
    "video_signal_segmentation",
    "live_fast",
    "spatial_effect_applied",
    "training_eligible",
    "review_required",
)


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text).replace("**", "").replace("`", "")


def document_text(document: Document) -> str:
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return f"{paragraph_text}\n{table_text}"


def find_missing_source_text(source_text: str, output_text: str) -> list[str]:
    missing: list[str] = []
    normalized_output = normalize(output_text)
    source_lines = source_text.splitlines()
    for index, line in enumerate(source_lines):
        stripped = line.strip()
        if not stripped or stripped.startswith(("!", "|", "---", "*图", "*Fig")):
            continue
        if stripped.startswith("# ") or stripped.startswith(("版本：", "日期：", "范围：", "项目定位：")):
            continue
        candidate = re.sub(r"^#{1,4}\s+", "", stripped)
        candidate = re.sub(r"^[-*]\s+", "", candidate)
        candidate = re.sub(r"^\d+(?:\.\d+)*\s+", "", candidate)
        if len(normalize(candidate)) < 12:
            continue
        candidate_words = normalize(candidate)
        if len(candidate_words) > 80:
            candidate_words = candidate_words[:80]
        if candidate_words not in normalized_output:
            missing.append(candidate)
    return missing


def validate_page(document: Document) -> list[str]:
    section = document.sections[0]
    errors: list[str] = []
    if abs(section.page_width.cm - 21) > 0.02 or abs(section.page_height.cm - 29.7) > 0.02:
        errors.append("纸张未设置为 A4")
    for label, value in (
        ("上页边距", section.top_margin.cm),
        ("下页边距", section.bottom_margin.cm),
        ("左页边距", section.left_margin.cm),
        ("右页边距", section.right_margin.cm),
    ):
        if abs(value - 2) > 0.02:
            errors.append(f"{label}不是 2 cm")
    grid = document.sections[0]._sectPr.find(qn("w:docGrid"))
    if grid is None or grid.get(qn("w:type")) != "linesAndChars":
        errors.append("文档网格不是 linesAndChars")
    return errors


def validate_spec(spec: DocumentSpec) -> list[str]:
    errors: list[str] = []
    if not spec.source.is_file():
        return [f"缺少源稿：{spec.source}"]
    if not spec.output.is_file():
        return [f"缺少 Word：{spec.output}"]
    source_text = spec.source.read_text(encoding="utf-8")
    document = Document(spec.output)
    output_text = document_text(document)
    errors.extend(validate_page(document))
    for term in spec.required_terms:
        if normalize(term) not in normalize(output_text):
            errors.append(f"缺少基线术语：{term}")
    for term in BANNED_INTERNAL_TERMS:
        if term in output_text:
            errors.append(f"包含内部命名：{term}")
    missing = find_missing_source_text(source_text, output_text)
    if missing:
        errors.append(f"正文未完整写入，缺少 {len(missing)} 行源稿内容：{missing[0]}")
    return errors


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data-summary-output", type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    specs = SPECS
    if args.data_summary_output is not None:
        specs = tuple(
            replace(spec, output=args.data_summary_output.resolve()) if spec.name == "数据汇总" else spec
            for spec in SPECS
        )
    failures: list[str] = []
    for spec in specs:
        errors = validate_spec(spec)
        if errors:
            failures.extend(f"{spec.name}：{error}" for error in errors)
            continue
        print(f"PASS {spec.name}：格式、基线术语、内部命名与正文一致性检查通过")
    if failures:
        for failure in failures:
            print(f"FAIL {failure}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
