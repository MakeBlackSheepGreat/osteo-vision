"""按 2020 版论文汇编格式生成数据汇总 Word 文档。

将 challenge_cup_report_draft_20260721/数据汇总.md 转为单一 docx，
严格应用：A4 页面、页边距 2cm、46字×43行网格、行距 16 磅；
题名二号黑体加粗居中、作者四号楷体_GB2312 居中；
标题层次 一、／（一）／1．；正文五号宋体；表格三线表小五号。
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============== 配置 ==============
ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "research" / "reports" / "submission" / "challenge_cup_report_draft_20260721"
INPUT_FILE = REPORT_DIR / "数据汇总.md"
OUTPUT_FILE = REPORT_DIR / "平台软件数据汇总.docx"

# 字体（缺失时 Word 会自动降级到系统等价字体）
SONG = "宋体"
HEI = "黑体"
KAI = "楷体_GB2312"
FANGSONG = "仿宋_GB2312"
TNR = "Times New Roman"

# 字号（pt）
PT_TITLE = 22  # 二号
PT_AUTHOR = 14  # 四号
PT_SUBHEAD = 12  # 小四
PT_BODY = 10.5  # 五号
PT_SMALL = 9  # 小五号

# 中文数字
CN_NUM = [
    "一",
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
    "十一",
    "十二",
    "十三",
    "十四",
    "十五",
]


# ============== 工具函数 ==============
def set_run_font(run, font_cn=SONG, font_en=TNR, size=PT_BODY, bold=False, italic=False):
    """设置 run 字体，中文需同时设 eastAsia。"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)


def set_line_spacing(p, line_pt=16, rule=WD_LINE_SPACING.EXACTLY):
    pf = p.paragraph_format
    pf.line_spacing_rule = rule
    pf.line_spacing = Pt(line_pt)


def set_first_indent_chars(p, chars=2, char_size_pt=PT_BODY):
    """首行缩进 N 字符。"""
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    ind.set(qn("w:firstLine"), str(int(chars * char_size_pt * 20)))


def add_body_text(doc, text, indent=2):
    """正文段落：五号宋体，首行缩进 2 字符。"""
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    set_first_indent_chars(p, indent, PT_BODY)
    run = p.add_run(text)
    set_run_font(run, SONG, TNR, PT_BODY)
    return p


def add_heading_l1(doc, cn_num, text):
    """一级标题：一、xxx，四号黑体加粗，段前段后 12pt。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 22)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    run = p.add_run(f"{cn_num}、{text}")
    set_run_font(run, HEI, TNR, PT_AUTHOR, bold=True)


def add_heading_l2(doc, cn_num, text):
    """二级标题：（一）xxx，小四黑体加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run(f"（{cn_num}）{text}")
    set_run_font(run, HEI, TNR, PT_SUBHEAD, bold=True)


def add_heading_l3(doc, num, text):
    """三级标题：1. xxx，小四宋体加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    run = p.add_run(f"{num}. {text}")
    set_run_font(run, SONG, TNR, PT_SUBHEAD, bold=True)


def add_list_item(doc, text, ordered=False, num_str="1"):
    """列表项：有序 1. xxx / 无序 • xxx。"""
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    set_first_indent_chars(p, 2, PT_BODY)
    if ordered:
        run = p.add_run(f"{num_str}. {text}")
    else:
        run = p.add_run(f"• {text}")
    set_run_font(run, SONG, TNR, PT_BODY)


def add_table_three_line(doc, header, rows, caption=None):
    """三线表：表头加粗，去左右边框，仅保留顶/底/表头下三条线。小五号。"""
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(p, 16)
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        run = p.add_run(caption)
        set_run_font(run, SONG, TNR, PT_SMALL, bold=True)

    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(p, 14)
        run = p.add_run(cell_text)
        set_run_font(run, HEI, TNR, PT_SMALL, bold=True)

    # 数据行
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(p, 14)
            run = p.add_run(cell_text)
            set_run_font(run, SONG, TNR, PT_SMALL)

    # 三线表样式：去左右边框，仅顶/底/表头下
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        if border_name in {"top", "bottom"}:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "12")
            border.set(qn("w:color"), "000000")
        else:
            border.set(qn("w:val"), "none")
        tblBorders.append(border)
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "single")
    insideH.set(qn("w:sz"), "4")
    insideH.set(qn("w:color"), "000000")
    tblBorders.append(insideH)
    insideV = OxmlElement("w:insideV")
    insideV.set(qn("w:val"), "none")
    tblBorders.append(insideV)
    tblPr.insert(2, tblBorders)

    # 表头行底边加粗
    for cell in table.rows[0].cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "000000")
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

    # 表后空行
    p = doc.add_paragraph()
    set_line_spacing(p, 10)
    return table


def add_horizontal_rule(doc):
    """添加分隔符（空段落）。"""
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)


# ============== Markdown 解析 ==============
def parse_markdown_to_docx(md_text, doc):
    """解析 markdown 文本并写入 docx。"""
    lines = md_text.splitlines()
    i = 0
    l1_counter = 0
    l2_counter = 0
    l3_counter = 0
    in_table = False
    table_header = []
    table_rows = []

    def flush_table():
        nonlocal table_header, table_rows, in_table
        if table_header:
            add_table_three_line(doc, table_header, table_rows)
        table_header = []
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 表格处理
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^:?-{2,}:?$", c) or c == "" for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线 ---
        if stripped == "---" or stripped == "***":
            add_horizontal_rule(doc)
            i += 1
            continue

        # 一级标题 # xxx
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title_text = stripped[2:].strip()
            l1_counter += 1
            l2_counter = 0
            l3_counter = 0
            add_heading_l1(doc, CN_NUM[l1_counter - 1], title_text)
            i += 1
            continue

        # 二级标题 ## xxx
        if stripped.startswith("## ") and not stripped.startswith("### "):
            title_text = stripped[3:].strip()
            # 去掉 "N.N " 前缀
            title_text = re.sub(r"^\d+(\.\d+)*\s*", "", title_text)
            l2_counter += 1
            l3_counter = 0
            add_heading_l2(doc, CN_NUM[l2_counter - 1], title_text)
            i += 1
            continue

        # 三级标题 ### xxx
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            title_text = stripped[4:].strip()
            title_text = re.sub(r"^\d+(\.\d+)*\s*", "", title_text)
            l3_counter += 1
            add_heading_l3(doc, l3_counter, title_text)
            i += 1
            continue

        # 引用块 >
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            if quote_text:
                p = doc.add_paragraph()
                set_line_spacing(p, 16)
                set_first_indent_chars(p, 2, PT_BODY)
                run = p.add_run(quote_text)
                set_run_font(run, KAI, TNR, PT_SMALL, italic=True)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 有序列表 1. xxx
        ol_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if ol_match:
            num = ol_match.group(1)
            text = ol_match.group(2)
            add_list_item(doc, text, ordered=True, num_str=num)
            i += 1
            continue

        # 无序列表 - xxx 或 * xxx
        if (stripped.startswith("- ") or stripped.startswith("* ")) and not stripped.startswith("**"):
            text = stripped[2:].strip()
            add_list_item(doc, text, ordered=False)
            i += 1
            continue

        # 图注 *xxx*（斜体说明）
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            caption_text = stripped.strip("*").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(p, 16)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(6)
            run = p.add_run(caption_text)
            set_run_font(run, SONG, TNR, PT_SMALL, italic=True)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 普通段落（合并连续行）
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith(">")
                or next_line.startswith("- ")
                or next_line.startswith("* ")
                or next_line == "---"
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            para_lines.append(next_line)
            j += 1
        text = " ".join(para_lines)
        if text:
            add_body_text(doc, text, indent=2)
        i = j

    # 末尾表格
    if in_table:
        flush_table()


# ============== 封面与摘要 ==============
def add_title_block(doc):
    """题名 + 作者 + 摘要 + 关键词。"""
    # 题名：二号黑体加粗居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 28)
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    run = p.add_run("平台软件数据汇总")
    set_run_font(run, HEI, TNR, PT_TITLE, bold=True)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_after = Pt(18)
    run = p.add_run("——研发验证数据与工程核验汇总")
    set_run_font(run, HEI, TNR, 16, bold=False)

    # 作者：四号楷体_GB2312 居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    run = p.add_run("项目团队")
    set_run_font(run, KAI, TNR, PT_AUTHOR, bold=False)

    # 摘要
    abstract_zh = (
        "本文档汇总平台软件当前的官方输入边界、数据分层、模型工程验证、4K 融合与视频处理性能、"
        "三维工程验证和质量核对结果。来源核验集合与分层数据注册表采用独立统计口径。"
        "全部模型和性能数据均来自公开、代理、近域或数字仿体工程资料；真实术中 ICG 颌骨骨髓炎"
        "记录与训练准入记录均为零。"
    )
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    set_first_indent_chars(p, 2, PT_SMALL)
    run = p.add_run("【摘要】")
    set_run_font(run, HEI, TNR, PT_SMALL, bold=True)
    run = p.add_run(abstract_zh)
    set_run_font(run, SONG, TNR, PT_SMALL)

    # 关键词
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    set_first_indent_chars(p, 2, PT_SMALL)
    pf = p.paragraph_format
    pf.space_after = Pt(12)
    run = p.add_run("【关键词】")
    set_run_font(run, HEI, TNR, PT_SMALL, bold=True)
    run = p.add_run("数据汇总；设备参数；模型性能；运行门控；工程验证")
    set_run_font(run, SONG, TNR, PT_SMALL)


# ============== 主流程 ==============
def setup_page(doc):
    """A4 页面、页边距 2cm、46字×43行网格。"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 46字×43行网格
    sectPr = section._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "linesAndChars")
    docGrid.set(qn("w:linePitch"), "312")
    docGrid.set(qn("w:charSpace"), "0")

    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")


def main():
    doc = Document()
    setup_page(doc)
    add_title_block(doc)

    # 读取数据汇总 md
    md_text = INPUT_FILE.read_text(encoding="utf-8")
    parse_markdown_to_docx(md_text, doc)

    output_file = OUTPUT_FILE
    if len(sys.argv) == 2:
        output_file = Path(sys.argv[1]).resolve()
    elif len(sys.argv) > 2:
        raise SystemExit("用法：python scripts/generate_data_summary_docx.py [输出 Word 路径]")

    temporary_output = output_file.with_name(f".{output_file.stem}.tmp{output_file.suffix}")
    try:
        doc.save(str(temporary_output))
        temporary_output.replace(output_file)
    finally:
        if temporary_output.exists():
            temporary_output.unlink()
    size_kb = output_file.stat().st_size / 1024
    print(f"[OK] 已生成：{output_file}")
    print(f"     文件大小：{size_kb:.1f} KB")


if __name__ == "__main__":
    main()
