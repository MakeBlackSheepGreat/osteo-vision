"""按 2020 版论文汇编格式生成 Word 文档。

将 challenge_cup_report_draft_20260721 目录下提纲与 Cap1-9 markdown 合并为单一 docx，
严格应用：A4 页面、页边距 2cm、46字×43行网格、行距 16 磅；
题名二号黑体加粗居中、作者四号楷体_GB2312 居中、摘要小五号宋体；
标题层次 一、／（一）／1．／（1）；正文五号宋体；表格三线表小五号。
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============== 配置 ==============
ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "research" / "reports" / "submission" / "challenge_cup_report_draft_20260721"
OUTPUT_FILE = REPORT_DIR / "颌骨骨髓炎智能化荧光诊疗比赛方案.docx"

# 字体（缺失时 Word 会自动降级到系统等价字体）
SONG = "宋体"
HEI = "黑体"
KAI = "楷体_GB2312"
FANGSONG = "仿宋_GB2312"
TNR = "Times New Roman"

# 字号（pt）
PT_TITLE = 22  # 二号
PT_AUTHOR = 14  # 四号
PT_SUBHEAD = 12  # 小四
PT_BODY = 10.5  # 五号
PT_SMALL = 9  # 小五号

# 章节文件顺序
CAP_FILES = [
    "Cap1_项目背景与临床需求.md",
    "Cap2_总体技术方案设计.md",
    "Cap3_荧光造影剂设计.md",
    "Cap4_多模态荧光显微图像融合方法.md",
    "Cap5_病灶智能识别与辅助决策模型.md",
    "Cap6_实验设计及结果验证.md",
    "Cap7_工程实现与临床应用方案.md",
    "Cap8_创新性与可行性.md",
    "Cap9_未来展望.md",
]

# 中文数字
CN_NUM = [
    "一",
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
    "十一",
    "十二",
    "十三",
    "十四",
    "十五",
]


# ============== 工具函数 ==============
def set_run_font(run, font_cn=SONG, font_en=TNR, size=PT_BODY, bold=False, italic=False):
    """设置 run 字体，中文需同时设 eastAsia。"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)


def set_line_spacing(p, line_pt=16, rule=WD_LINE_SPACING.EXACTLY):
    pf = p.paragraph_format
    pf.line_spacing_rule = rule
    pf.line_spacing = Pt(line_pt)


def set_first_indent_chars(p, chars=2, char_size_pt=PT_BODY):
    """首行缩进 N 字符。"""
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    # 使用 firstLineChars（按字符数），同时给 firstLine 兜底
    ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
    # firstLine 单位 twips，1pt=20twips
    ind.set(qn("w:firstLine"), str(int(chars * char_size_pt * 20)))


def setup_page(doc):
    """A4 + 页边距 2cm + 46字×43行网格 + 行距 16 磅。"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 文档网格：46 字 × 43 行
    sectPr = section._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "linesAndChars")
    docGrid.set(qn("w:linePitch"), "312")  # ≈15.6pt，与 16 磅接近
    docGrid.set(qn("w:charSpace"), "0")


def add_para(
    doc,
    text="",
    font_cn=SONG,
    font_en=TNR,
    size=PT_BODY,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_pt=16,
    indent_chars=0,
    space_before=0,
    space_after=0,
):
    """添加普通段落。"""
    p = doc.add_paragraph()
    p.alignment = align
    set_line_spacing(p, line_pt)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_chars > 0:
        set_first_indent_chars(p, indent_chars, size)
    if text:
        runs = parse_inline_emphasis(text)
        for rtext, rbold in runs:
            run = p.add_run(rtext)
            set_run_font(run, font_cn, font_en, size, bold=bold or rbold)
    return p


def parse_inline_emphasis(text):
    """解析 **加粗**，返回 [(text, bold)] 列表。"""
    parts = []
    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos : m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    return parts if parts else [(text, False)]


def add_heading_l1(doc, cn_num, text):
    """一级标题：一、xxx —— 四号黑体居中。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = p.add_run(f"{cn_num}、{text}")
    set_run_font(run, HEI, TNR, PT_AUTHOR, bold=True)  # 四号 14pt
    return p


def add_heading_l2(doc, cn_num, text):
    """二级标题：（一）xxx —— 仿宋_GB2312 小四加黑空两格。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    set_first_indent_chars(p, 2, PT_SUBHEAD)
    run = p.add_run(f"（{cn_num}）{text}")
    set_run_font(run, FANGSONG, TNR, PT_SUBHEAD, bold=True)
    return p


def add_heading_l3(doc, num, text):
    """三级标题：1．xxx —— 五号空两格。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(2)
    set_first_indent_chars(p, 2, PT_BODY)
    run = p.add_run(f"{num}．{text}")
    set_run_font(run, SONG, TNR, PT_BODY, bold=True)
    return p


def add_heading_l4(doc, num, text):
    """四级标题：（1）xxx —— 五号空两格。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(2)
    set_first_indent_chars(p, 2, PT_BODY)
    run = p.add_run(f"（{num}）{text}")
    set_run_font(run, SONG, TNR, PT_BODY, bold=True)
    return p


def add_body_text(doc, text, indent=2):
    """正文段落：五号宋体，首行空二格。"""
    return add_para(
        doc,
        text,
        font_cn=SONG,
        font_en=TNR,
        size=PT_BODY,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_pt=16,
        indent_chars=indent,
        space_before=0,
        space_after=0,
    )


def add_quote_block(doc, text):
    """引用块：小五号宋体，左缩进，斜体淡色（提示性）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    runs = parse_inline_emphasis(text)
    for rtext, rbold in runs:
        run = p.add_run(rtext)
        set_run_font(run, SONG, TNR, PT_SMALL, bold=rbold, italic=False)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_list_item(doc, text, ordered=False, num_str=None):
    """列表项：五号宋体，首行缩进 2 字符。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.74)  # 约 2 字符
    pf.first_line_indent = Cm(-0.37)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    prefix = f"{num_str}. " if ordered else "• "
    run = p.add_run(prefix)
    set_run_font(run, SONG, TNR, PT_BODY)
    runs = parse_inline_emphasis(text)
    for rtext, rbold in runs:
        run = p.add_run(rtext)
        set_run_font(run, SONG, TNR, PT_BODY, bold=rbold)
    return p


def add_image_caption(doc, alt_text, img_path):
    """Insert a required report image with a centered caption."""
    image_path = Path(img_path)
    if not image_path.is_file():
        raise FileNotFoundError(f"Markdown image is missing from the report package: {image_path}")

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p_img, 16)
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Cm(12))
    # 图序图名（图下居中，小五号）
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p_cap, 16)
    pf = p_cap.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    # 解析 alt_text 形如 "Fig 1 xxx"
    caption = alt_text.replace("Fig", "图")
    caption = re.sub(r"图\s*(\d+)", r"图\1", caption)
    run = p_cap.add_run(caption)
    set_run_font(run, SONG, TNR, PT_SMALL, bold=False)


def set_cell_border(cell, top=True, bottom=True, left=False, right=False, inside_h=False, inside_v=False):
    """设置单元格边框（三线表用）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, flag in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
        ("insideH", inside_h),
        ("insideV", inside_v),
    ]:
        elem = tcBorders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tcBorders.append(elem)
        if flag:
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "6")  # 0.75pt
            elem.set(qn("w:color"), "000000")
        else:
            elem.set(qn("w:val"), "nil")


def add_table_three_line(doc, header, rows, caption_above=None):
    """三线表：表序表名在表上居中，小五号；表格去左右边框。"""
    # 表序表名
    if caption_above:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(p_cap, 16)
        pf = p_cap.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(0)
        run = p_cap.add_run(caption_above)
        set_run_font(run, SONG, TNR, PT_SMALL, bold=False)

    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(p, 16)
        run = p.add_run(h.strip())
        set_run_font(run, SONG, TNR, PT_SMALL, bold=True)
        # 表头：上边框 + 下边框
        set_cell_border(
            hdr_cells[i],
            top=True,
            bottom=True,
            left=False,
            right=False,
            inside_h=True,
            inside_v=False,
        )

    # 数据行
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            if c_idx >= n_cols:
                break
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(p, 16)
            run = p.add_run(val.strip())
            set_run_font(run, SONG, TNR, PT_SMALL, bold=False)
            is_bottom_row = r_idx == len(rows) - 1
            set_cell_border(
                cells[c_idx],
                top=False,
                bottom=is_bottom_row,
                left=False,
                right=False,
                inside_h=not is_bottom_row,
                inside_v=False,
            )

    # 表后空一行
    p_empty = doc.add_paragraph()
    set_line_spacing(p_empty, 16)
    pf = p_empty.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    return table


def add_horizontal_rule(doc):
    """添加分隔符（空段落）。"""
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)


def resolve_report_image_path(raw_path: str) -> Path:
    """Resolve a Markdown image without allowing it to leave the report package."""
    candidate = Path(raw_path)
    if candidate.is_absolute():
        raise ValueError(f"Markdown image must use a report-relative path: {raw_path}")

    report_root = REPORT_DIR.resolve()
    image_path = (report_root / candidate).resolve(strict=False)
    try:
        relative_path = image_path.relative_to(report_root)
    except ValueError as exc:
        raise ValueError(f"Markdown image path escapes the report package: {raw_path}") from exc
    if not relative_path.parts:
        raise ValueError(f"Markdown image path must name a file: {raw_path}")
    return image_path


# ============== Markdown 解析 ==============
def parse_markdown_to_docx(md_text, doc, cap_index):
    """解析 markdown 文本并写入 docx。"""
    lines = md_text.splitlines()
    i = 0
    # 二级/三级/四级标题计数器
    l2_counter = 0
    l3_counter = 0
    l4_counter = 0
    in_table = False
    table_header = []
    table_rows = []
    table_caption = None

    def flush_table():
        nonlocal table_header, table_rows, table_caption, in_table
        if table_header:
            add_table_three_line(doc, table_header, table_rows, table_caption)
        table_header = []
        table_rows = []
        table_caption = None
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 表格处理
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # 分隔行
            if all(re.match(r"^:?-{2,}:?$", c) or c == "" for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线 ---
        if stripped == "---" or stripped == "***":
            add_horizontal_rule(doc)
            i += 1
            continue

        # 一级标题 # Cap X xxx
        if stripped.startswith("# ") and not stripped.startswith("## "):
            # Cap 主标题转为一级（章）
            title_text = stripped[2:].strip()
            # 去掉 "Cap N " 前缀
            title_text = re.sub(r"^Cap\s*\d+\s*", "", title_text)
            add_heading_l1(doc, CN_NUM[cap_index], title_text)
            l2_counter = 0
            l3_counter = 0
            l4_counter = 0
            i += 1
            continue

        # 二级标题 ## N.N xxx
        if stripped.startswith("## ") and not stripped.startswith("### "):
            title_text = stripped[3:].strip()
            # 去掉 "N.N " 前缀
            title_text = re.sub(r"^\d+(\.\d+)*\s*", "", title_text)
            l2_counter += 1
            l3_counter = 0
            l4_counter = 0
            add_heading_l2(doc, CN_NUM[l2_counter - 1], title_text)
            i += 1
            continue

        # 三级标题 ### N.N.N xxx
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            title_text = stripped[4:].strip()
            title_text = re.sub(r"^\d+(\.\d+)*\s*", "", title_text)
            l3_counter += 1
            l4_counter = 0
            add_heading_l3(doc, l3_counter, title_text)
            i += 1
            continue

        # 四级标题 #### N.N.N.N xxx
        if stripped.startswith("#### "):
            title_text = stripped[5:].strip()
            title_text = re.sub(r"^\d+(\.\d+)*\s*", "", title_text)
            l4_counter += 1
            add_heading_l4(doc, l4_counter, title_text)
            i += 1
            continue

        # 引用块 >
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            if quote_text:
                add_quote_block(doc, quote_text)
            i += 1
            continue

        # 图片 ![alt](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = resolve_report_image_path(img_match.group(2))
            add_image_caption(doc, alt_text, str(img_path))
            i += 1
            continue

        # 有序列表 1. xxx
        ol_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if ol_match:
            num = ol_match.group(1)
            text = ol_match.group(2)
            add_list_item(doc, text, ordered=True, num_str=num)
            i += 1
            continue

        # 无序列表 - xxx 或 * xxx
        if (stripped.startswith("- ") or stripped.startswith("* ")) and not stripped.startswith("**"):
            text = stripped[2:].strip()
            add_list_item(doc, text, ordered=False)
            i += 1
            continue

        # 图注 *xxx*（斜体说明）
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            caption_text = stripped.strip("*").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(p, 16)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(6)
            run = p.add_run(caption_text)
            set_run_font(run, SONG, TNR, PT_SMALL, italic=True)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 普通段落
        # 合并连续段落
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith(">")
                or next_line.startswith("- ")
                or next_line.startswith("* ")
                or next_line.startswith("1.")
                or next_line == "---"
                or re.match(r"^!\[", next_line)
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            para_lines.append(next_line)
            j += 1
        text = " ".join(para_lines)
        if text:
            add_body_text(doc, text, indent=2)
        i = j

    # 末尾表格
    if in_table:
        flush_table()


# ============== 封面与摘要 ==============
def add_title_block(doc):
    """题名 + 作者 + 摘要 + 关键词。"""
    # 题名：二号黑体加粗居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 28)
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    run = p.add_run("颌骨骨髓炎智能化荧光诊疗比赛方案")
    set_run_font(run, HEI, TNR, PT_TITLE, bold=True)

    # 副标题（项目基线）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_after = Pt(18)
    run = p.add_run("——研发验证版平台软件工程方案")
    set_run_font(run, HEI, TNR, 16, bold=False)

    # 作者：四号楷体_GB2312 居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    run = p.add_run("项目团队")
    set_run_font(run, KAI, TNR, PT_AUTHOR, bold=False)

    # 单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 18)
    pf = p.paragraph_format
    pf.space_after = Pt(18)
    run = p.add_run("（osteo-vision 项目组）")
    set_run_font(run, KAI, TNR, PT_SMALL, bold=False)

    # 摘要标签 + 内容
    abstract_zh = (
        "本方案面向颌骨骨髓炎术中辅助决策需求，围绕赛题方荧光手术显微镜平台构建"
        "“新型荧光造影剂设计 + 多模态医学图像融合与处理 + 人工智能辅助显微成像判读”"
        "的集成方案。平台软件以 4K MP4/JPEG 为一级输入，承担荧光分析、AI 与医生交互判读、"
        "结果输出三层任务：基于 ICG 灌注层并补充四环素类骨活性文献证据与模块化新型骨亲和探针设计；"
        "采用白光与 ICG 双通道配准、伪彩、融合、归一化、质控与 ROI 定量；"
        "以 keyframe_residual_attention_unet_s 为主线模型输出骨面、荧光信号、风险、不确定性四类 mask "
        "与骨活性连续评分；通过医生复核回灌、证据包导出与 L1/L2 三维工程验证形成端到端闭环。"
        "所有输出属于研发验证证据，医生保留最终判断。"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    set_first_indent_chars(p, 2, PT_SMALL)
    run = p.add_run("摘要：")
    set_run_font(run, SONG, TNR, PT_SMALL, bold=True)
    run = p.add_run(abstract_zh)
    set_run_font(run, SONG, TNR, PT_SMALL, bold=False)

    # 关键词
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p, 16)
    pf = p.paragraph_format
    pf.space_after = Pt(12)
    set_first_indent_chars(p, 2, PT_SMALL)
    run = p.add_run("关键词：")
    set_run_font(run, SONG, TNR, PT_SMALL, bold=True)
    run = p.add_run("颌骨骨髓炎；ICG 荧光成像；多模态融合；AI 辅助判读；医生复核")
    set_run_font(run, SONG, TNR, PT_SMALL, bold=False)


# ============== 主流程 ==============
def main():
    cap_paths = [REPORT_DIR / cap_file for cap_file in CAP_FILES]
    missing_paths = [path for path in cap_paths if not path.is_file()]
    if missing_paths:
        missing_text = "\n".join(f"- {path}" for path in missing_paths)
        raise FileNotFoundError(f"Challenge Cup report is missing required chapter files:\n{missing_text}")

    doc = Document()
    setup_page(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = TNR
    style.font.size = Pt(PT_BODY)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), SONG)

    # 题名 + 作者 + 摘要 + 关键词
    add_title_block(doc)

    # 章节分隔
    add_horizontal_rule(doc)

    # 按 Cap 顺序处理
    for idx, cap_path in enumerate(cap_paths):
        cap_file = cap_path.name
        md_text = cap_path.read_text(encoding="utf-8")
        print(f"[INFO] 处理：{cap_file}（{len(md_text)} 字符）")
        parse_markdown_to_docx(md_text, doc, idx)

    temporary_output = OUTPUT_FILE.with_name(f".{OUTPUT_FILE.stem}.tmp{OUTPUT_FILE.suffix}")
    try:
        doc.save(str(temporary_output))
        temporary_output.replace(OUTPUT_FILE)
    finally:
        if temporary_output.exists():
            temporary_output.unlink()
    print(f"\n[OK] 已生成：{OUTPUT_FILE}")
    print(f"     文件大小：{OUTPUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
