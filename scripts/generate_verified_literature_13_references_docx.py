from __future__ import annotations

import csv
import json
import re
import sys
import urllib.error
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
MANIFEST = ROOT / "research" / "literature" / "inventory" / "literature_13_audit_manifest_20260728.csv"
OUTPUT_DIR = ROOT / "research" / "reports" / "submission"
OUTPUT_DOCX = OUTPUT_DIR / "osteo_vision_verified_literature_13_gbt7714_2015_20260728.docx"
OUTPUT_AUDIT = OUTPUT_DIR / "osteo_vision_verified_literature_13_gbt7714_2015_20260728_validation.md"


def normalize(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def fetch_crossref(doi: str) -> dict:
    request = urllib.request.Request(
        f"https://api.crossref.org/works/{urllib.parse.quote(doi, safe='')}",
        headers={"User-Agent": "osteo-vision-reference-builder/1.0 (metadata verification)"},
    )
    with urllib.request.urlopen(request, timeout=30) as response:
        payload = json.load(response)
    return payload["message"]


def publication_year(record: dict) -> int:
    for field in ("published-print", "published-online", "issued", "created"):
        parts = record.get(field, {}).get("date-parts", [[]])
        if parts and parts[0]:
            return int(parts[0][0])
    raise ValueError("Crossref metadata does not contain a publication year")


def authors(record: dict) -> str:
    names = record.get("author", [])
    if not names:
        raise ValueError("Crossref metadata does not contain authors")
    formatted = []
    for author in names[:3]:
        family = author.get("family", "").strip()
        given = author.get("given", "").strip()
        initials = "".join(part[0].upper() for part in re.findall(r"[A-Za-z]+", given) if part)
        formatted.append(f"{family} {initials}".strip())
    if len(names) > 3:
        formatted.append("et al")
    return ", ".join(formatted)


def source(record: dict) -> str:
    containers = record.get("container-title", [])
    if not containers:
        raise ValueError("Crossref metadata does not contain a journal title")
    return containers[0].strip()


def location(record: dict) -> str:
    volume = str(record.get("volume", "")).strip()
    issue = str(record.get("issue", "")).strip()
    pages = str(record.get("page", record.get("article-number", ""))).strip()
    if not volume:
        return pages
    if issue:
        volume = f"{volume}({issue})"
    return f"{volume}: {pages}" if pages else volume


def citation(record: dict, index: int) -> str:
    title = record.get("title", [""])[0].strip().rstrip(".")
    if not title:
        raise ValueError("Crossref metadata does not contain a title")
    doi = record["DOI"].lower()
    return f"[{index}] {authors(record)}. {title}[J]. {source(record)}, {publication_year(record)}, {location(record)}. DOI: {doi}."


def set_font(run, chinese_font: str, latin_font: str, size: float, bold: bool = False) -> None:
    run.font.name = latin_font
    run.font.size = Pt(size)
    run.font.bold = bold
    properties = run._element.get_or_add_rPr()
    fonts = properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.append(fonts)
    fonts.set(qn("w:eastAsia"), chinese_font)
    fonts.set(qn("w:ascii"), latin_font)
    fonts.set(qn("w:hAnsi"), latin_font)


def set_paragraph_layout(paragraph, hanging: bool = False) -> None:
    format_ = paragraph.paragraph_format
    format_.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    format_.line_spacing = Pt(20)
    format_.space_before = Pt(0)
    format_.space_after = Pt(0)
    if hanging:
        format_.left_indent = Cm(0.74)
        format_.first_line_indent = Cm(-0.74)


def build_document(citations: list[str]) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    normal_fonts = style._element.rPr.rFonts
    normal_fonts.set(qn("w:eastAsia"), "宋体")
    normal_fonts.set(qn("w:ascii"), "Times New Roman")
    normal_fonts.set(qn("w:hAnsi"), "Times New Roman")

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(12)
    heading_run = heading.add_run("参考文献")
    set_font(heading_run, "黑体", "Times New Roman", 16, bold=True)

    for item in citations:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_layout(paragraph, hanging=True)
        set_font(paragraph.add_run(item), "宋体", "Times New Roman", 10.5)

    document.core_properties.title = "已核验13篇文献参考文献"
    document.core_properties.subject = "GB/T 7714-2015 顺序编码制"
    document.core_properties.author = "osteo-vision"
    document.core_properties.comments = "依据本地 literature_13_audit_manifest_20260728.csv 及 Crossref 元数据生成。"
    document.save(OUTPUT_DOCX)


def write_audit(rows: list[dict], records: list[dict], citations: list[str]) -> None:
    lines = [
        "# GB/T 7714-2015 参考文献生成核验记录",
        "",
        "- 生成日期：2026-07-28",
        "- 来源清单：`research/literature/inventory/literature_13_audit_manifest_20260728.csv`",
        "- 元数据复核：逐条调用 Crossref Works，核对 DOI 与英文题名。",
        "- 编排规则：GB/T 7714-2015 顺序编码制；外文期刊文献使用 `[J]` 标识，作者列至前三位后使用 `et al`。",
        "- Word 版式：A4；标题黑体小二；条目宋体/Times New Roman 五号；固定 20 磅行距；0.74 cm 悬挂缩进。",
        "",
        "## 核验结果",
        "",
        "| 编号 | DOI | 题名匹配 | Crossref 年份 | 引文状态 |",
        "| --- | --- | --- | ---: | --- |",
    ]
    for row, record, item in zip(rows, records, citations, strict=True):
        matches = normalize(row["english_title"]) == normalize(record["title"][0])
        doi_matches = row["doi"].lower() == record["DOI"].lower()
        status = "通过" if matches and doi_matches and item.startswith(f"[{int(row['audit_id'][1:]):d}]") else "不通过"
        lines.append(
            f"| {row['audit_id']} | {row['doi']} | {'是' if matches and doi_matches else '否'} | {publication_year(record)} | {status} |"
        )
    lines.extend(
        [
            "",
            "## 结论",
            "",
            f"共生成 {len(citations)} 条参考文献；13 条 DOI 均与已核验清单及 Crossref Works 返回值一致，题名均匹配。",
            "",
        ]
    )
    OUTPUT_AUDIT.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    with MANIFEST.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if len(rows) != 13:
        raise ValueError(f"Expected 13 literature records, found {len(rows)}")
    records = []
    citations = []
    for index, row in enumerate(rows, start=1):
        record = fetch_crossref(row["doi"])
        if record.get("DOI", "").lower() != row["doi"].lower():
            raise ValueError(f"DOI mismatch for {row['audit_id']}")
        if normalize(record.get("title", [""])[0]) != normalize(row["english_title"]):
            raise ValueError(f"Title mismatch for {row['audit_id']}")
        records.append(record)
        citations.append(citation(record, index))
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_document(citations)
    write_audit(rows, records, citations)
    print(OUTPUT_DOCX)
    print(OUTPUT_AUDIT)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (OSError, ValueError, urllib.error.URLError, KeyError) as error:
        print(f"Generation failed: {error}", file=sys.stderr)
        raise SystemExit(1)
